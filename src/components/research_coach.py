"""Research Coach Component"""

import streamlit as st
from datetime import datetime
from src.database.database import get_db_session
from src.database.models import ResearchConversation
from src.services.gemini_service import GeminiService
try:
    from docx import Document
except ImportError:
    Document = None  # Fallback if python-docx not available
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io
from src.components.ui.card import card
from src.utils.helpers import format_date_local


def export_to_docx(content: str, filename: str = "research_draft.docx"):
    """Export content to DOCX"""
    if Document is None:
        raise ImportError("python-docx is required for DOCX export")
    doc = Document()
    doc.add_heading('Research Document', 0)
    doc.add_paragraph(content)
    return doc


def export_to_pdf(content: str, filename: str = "research_draft.pdf"):
    """Export content to PDF"""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    y = height - 50
    lines = content.split('\n')
    for line in lines:
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(50, y, line[:80])
        y -= 20
    
    c.save()
    buffer.seek(0)
    return buffer.getvalue()


def render_research_coach():
    """Render research coach interface"""
    st.title("💬 Research Coach")
    st.markdown("Brainstorm ideas, refine your thinking, and generate structured research documents.")
    
    db = get_db_session()
    user_id = st.session_state.user_id
    gemini_service: GeminiService = st.session_state.gemini_service
    
    if not gemini_service.is_configured():
        st.error("Please configure your Gemini API key first.")
        return
    
    try:
        # Initialize session state
        if 'research_conversation' not in st.session_state:
            st.session_state.research_conversation = []
        
        if 'research_outline' not in st.session_state:
            st.session_state.research_outline = ""
        
        if 'research_draft' not in st.session_state:
            st.session_state.research_draft = ""
        
        # Input topic
        topic = st.text_input("Research Topic", placeholder="e.g., The impact of AI on education")
        
        # Main workflow
        tab1, tab2, tab3, tab4 = st.tabs(["💭 Brainstorm", "📝 Outline", "✍️ Draft", "📚 History"])
        
        with tab1:
            st.markdown("### Start Your Research")
            initial_thoughts = st.text_area(
                "Your Initial Thoughts",
                placeholder="What are your initial ideas about this topic? What questions do you have?",
                height=150
            )
            
            if st.button("🤔 Get Research Questions"):
                if topic and initial_thoughts:
                    with st.spinner("Analyzing your thoughts..."):
                        result = gemini_service.brainstorm_research(topic, initial_thoughts)
                        
                        st.markdown("### Questions to Consider")
                        for q in result.get('questions', []):
                            card("", q)
                        
                        if result.get('suggestions'):
                            st.markdown("### Suggested Angles")
                            for suggestion in result.get('suggestions', []):
                                st.markdown(f"- {suggestion}")
                        
                        if result.get('gaps'):
                            st.markdown("### Areas Needing Clarity")
                            for gap in result.get('gaps', []):
                                st.markdown(f"- {gap}")
                        
                        refined_focus = result.get('refined_focus', topic)
                        st.session_state.research_outline = ""
                        st.session_state.research_draft = ""
                        
                        st.info(f"**Refined Focus:** {refined_focus}")
                
                else:
                    st.warning("Please enter both topic and initial thoughts.")
        
        with tab2:
            st.markdown("### Generate Outline")
            
            if topic:
                key_points = st.text_area(
                    "Key Points to Include (one per line)",
                    placeholder="Point 1\nPoint 2\nPoint 3",
                    height=150
                )
                
                if st.button("📋 Generate Outline"):
                    points_list = [p.strip() for p in key_points.split('\n') if p.strip()]
                    research_focus = topic  # Could use refined focus from brainstorm
                    
                    with st.spinner("Creating outline..."):
                        outline = gemini_service.generate_outline(topic, research_focus, points_list)
                        st.session_state.research_outline = outline
                    
                    card("Research Outline", f"<pre style='white-space: pre-wrap;'>{outline}</pre>")
                else:
                    if st.session_state.research_outline:
                        card("Research Outline", f"<pre style='white-space: pre-wrap;'>{st.session_state.research_outline}</pre>")
            else:
                st.info("Please enter a research topic first.")
        
        with tab3:
            st.markdown("### Generate Draft")
            
            if st.session_state.research_outline:
                if st.button("✍️ Generate Draft"):
                    with st.spinner("Writing draft..."):
                        draft = gemini_service.generate_draft(st.session_state.research_outline, topic)
                        st.session_state.research_draft = draft
                    
                    card("Research Draft", f"<div style='white-space: pre-wrap;'>{draft}</div>")
                    
                    # Export options
                    st.markdown("### Export Options")
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # DOCX
                        doc = export_to_docx(st.session_state.research_draft)
                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        st.download_button(
                            "📄 Download DOCX",
                            buffer.getvalue(),
                            file_name="research_draft.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    with col2:
                        # PDF
                        pdf_bytes = export_to_pdf(st.session_state.research_draft)
                        st.download_button(
                            "📑 Download PDF",
                            pdf_bytes,
                            file_name="research_draft.pdf",
                            mime="application/pdf"
                        )
                    
                    with col3:
                        # Markdown
                        st.download_button(
                            "📝 Download MD",
                            st.session_state.research_draft.encode(),
                            file_name="research_draft.md",
                            mime="text/markdown"
                        )
                    
                    with col4:
                        # TXT
                        st.download_button(
                            "📄 Download TXT",
                            st.session_state.research_draft.encode(),
                            file_name="research_draft.txt",
                            mime="text/plain"
                        )
                else:
                    if st.session_state.research_draft:
                        card("Research Draft", f"<div style='white-space: pre-wrap;'>{st.session_state.research_draft}</div>")
            else:
                st.info("Please generate an outline first.")
        
        with tab4:
            st.markdown("### Conversation History")
            conversations = db.query(ResearchConversation).filter(
                ResearchConversation.user_id == user_id
            ).order_by(ResearchConversation.created_at.desc()).limit(10).all()
            
            if conversations:
                for conv in conversations:
                    card(
                        f"{conv.title or 'Untitled'} - {format_date_local(conv.created_at)}",
                        f"""
                        <details>
                        <summary>View conversation</summary>
                        <pre>{str(conv.messages)}</pre>
                        </details>
                        """
                    )
            else:
                st.info("No conversations yet.")
            
            # Save current conversation
            if st.session_state.research_conversation and topic:
                if st.button("💾 Save Conversation"):
                    conv = ResearchConversation(
                        user_id=user_id,
                        title=topic,
                        messages=st.session_state.research_conversation,
                        outline=st.session_state.research_outline,
                        draft=st.session_state.research_draft
                    )
                    db.add(conv)
                    db.commit()
                    st.success("Conversation saved!")
    
    finally:
        db.close()

